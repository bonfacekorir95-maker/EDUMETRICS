
# Edumetrics main (compact)
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import pandas as pd, os
from docx import Document
from docx.shared import Pt

APP_TITLE = "Edumetrics 1.0 - Elgeyo Sawmill Edition"
SCHOOL = "ELGEYO SAWMILL PRIMARY AND JUNIOR SCHOOL"
SUBJECTS = ["Mathematics","English","Kiswahili","Integrated Science","Agriculture","Social Studies","CRE","Creative Arts & Sports","Pretechnical Studies"]

def perf(avg):
    if avg>=87.5: return "EE1"
    if avg>=75: return "EE2"
    if avg>=62.5: return "ME1"
    if avg>=50: return "ME2"
    if avg>=37.5: return "AE1"
    if avg>=25: return "AE2"
    if avg>=12.5: return "BE1"
    return "BE2"

def compute(df):
    for s in SUBJECTS:
        if s not in df.columns: df[s]=0
    df['Total']=df[SUBJECTS].sum(axis=1)
    df['Average']=(df['Total']/len(SUBJECTS)).round(2)
    df['Performance']=df['Average'].apply(perf)
    df['Pos']=df['Total'].rank(method='min', ascending=False).astype(int)
    return df

class App(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title(APP_TITLE); self.state('zoomed')
        men = tk.Menu(self); filem=tk.Menu(men,tearoff=0); filem.add_command(label="Open sample",command=self.load_sample); filem.add_command(label="Exit",command=self.quit); men.add_cascade(label="File",menu=filem)
        men.add_cascade(label="Registration", menu=tk.Menu(men))
        men.add_cascade(label="Administration", menu=tk.Menu(men))
        men.add_cascade(label="Records", menu=tk.Menu(men))
        men.add_cascade(label="Academics", menu=tk.Menu(men))
        men.add_cascade(label="SMS/Messaging", menu=tk.Menu(men))
        self.config(menu=men)
        self.tree=ttk.Treeview(self, columns=('Adm','Name','Class','Total','Average','Performance'), show='headings')
        for c in ('Adm','Name','Class','Total','Average','Performance'): self.tree.heading(c,text=c)
        self.tree.pack(fill='both', expand=True)
        btnf=tk.Frame(self); btnf.pack(pady=6)
        tk.Button(btnf,text='Load sample', command=self.load_sample).pack(side='left', padx=6)
        tk.Button(btnf,text='Generate reports', command=self.generate_reports).pack(side='left', padx=6)
        self.df=None
    def load_sample(self):
        p=os.path.join(os.getcwd(),'sample_results.xlsx')
        if not os.path.exists(p): messagebox.showerror('Missing','Place sample_results.xlsx in this folder'); return
        df=pd.read_excel(p); df=compute(df); self.df=df; self.populate()
    def populate(self):
        for i in self.tree.get_children(): self.tree.delete(i)
        for _,r in self.df.iterrows():
            self.tree.insert('', 'end', values=(r.get('AdmNo',''), r.get('Name',''), r.get('Class',''), r.get('Total',''), r.get('Average',''), r.get('Performance','')))
    def generate_reports(self):
        if self.df is None: messagebox.showerror('No data','Load sample first'); return
        out=os.path.join(os.getcwd(),'output'); os.makedirs(out, exist_ok=True)
        for _,r in self.df.iterrows():
            doc=Document(); h=doc.add_heading(level=1); h.alignment=1; run=h.add_run(SCHOOL); run.bold=True; run.font.size=Pt(16)
            doc.add_paragraph('Report Card'); doc.add_paragraph(f"Name: {r.get('Name')}  AdmNo: {r.get('AdmNo')}")
            tbl=doc.add_table(rows=1,cols=4); hdr=tbl.rows[0].cells; hdr[0].text='Subject'; hdr[1].text='Marks'; hdr[2].text='Grade'; hdr[3].text='Performance'
            for s in SUBJECTS:
                row=tbl.add_row().cells; row[0].text=s; row[1].text=str(r.get(s,0)); row[2].text=''; row[3].text=str(r.get('Performance',''))
            fname=os.path.join(out, f"{r.get('AdmNo')}_{r.get('Name')}_report.docx".replace(' ','_')); doc.save(fname)
        messagebox.showinfo('Done','Reports saved to output folder')

if __name__=='__main__':
    app=App(); app.mainloop()
